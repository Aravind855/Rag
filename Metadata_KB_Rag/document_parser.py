"""
Document Parser & Structure-Aware Extraction Module for NexaCore Knowledge Base RAG (Project 2).

Defines custom LlamaIndex `BaseReader` classes to parse PDF and DOCX documents with
full structure preservation (converting tables into Markdown grid tables `| ... |`,
preserving section headings `#`, `##`, and bullet points). Also enriches Document
instances with domain classification metadata.
"""

import logging
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from llama_index.core import Document
    from llama_index.core.readers.base import BaseReader
except ImportError:
    class Document:  # type: ignore
        def __init__(self, text: str = "", metadata: Optional[Dict[str, Any]] = None, **kwargs: Any) -> None:
            self.text = text
            self.metadata = metadata or {}
            self.doc_id = kwargs.get("doc_id", "")

    class BaseReader:  # type: ignore
        pass

# PyMuPDF / PyMuPDF4LLM for advanced PDF structure & table parsing
try:
    import fitz  # PyMuPDF
    import pymupdf4llm
except ImportError:
    fitz = None
    pymupdf4llm = None

# python-docx for DOCX paragraph and table parsing
try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger(__name__)


class StructureAwarePDFReader(BaseReader):
    """Structure-aware PDF Reader for LlamaIndex preserving tables, headers, and bullet points."""

    def __init__(self, use_pymupdf4llm: bool = True):
        super().__init__()
        self.use_pymupdf4llm = use_pymupdf4llm

    def load_data(
        self,
        file: Path,
        extra_info: Optional[Dict[str, Any]] = None,
        **load_kwargs: Any,
    ) -> List[Document]:
        """Parse PDF file into a LlamaIndex Document object with table grid preservation."""
        file_path = Path(file).resolve()
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        metadata = dict(extra_info) if extra_info else {}
        total_pages = 1

        if fitz is not None:
            try:
                with fitz.open(file_path) as pdf_doc:
                    total_pages = pdf_doc.page_count
            except Exception as e:
                logger.warning(f"Could not read page count for '{file_path.name}': {e}")

        metadata["total_pages"] = total_pages

        # 1. Primary: Use PyMuPDF4LLM for table and markdown structure preservation
        if self.use_pymupdf4llm and pymupdf4llm is not None:
            try:
                md_text = pymupdf4llm.to_markdown(str(file_path))
                if md_text and md_text.strip():
                    doc = Document(
                        text=md_text.strip(),
                        metadata=metadata,
                    )
                    logger.info(
                        f"Parsed PDF '{file_path.name}' via PyMuPDF4LLM: "
                        f"{len(md_text):,} chars across {total_pages} page(s) (Tables preserved)."
                    )
                    return [doc]
            except Exception as e:
                logger.warning(f"PyMuPDF4LLM extraction failed for '{file_path.name}', falling back: {e}")

        # 2. Fallback: Standard PyMuPDF page text extraction
        if fitz is not None:
            try:
                with fitz.open(file_path) as pdf_doc:
                    pages_text = []
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc[page_num]
                        page_str = page.get_text()
                        if page_str.strip():
                            pages_text.append(f"<!-- Page {page_num + 1} -->\n{page_str.strip()}")
                    full_text = "\n\n".join(pages_text)
                    doc = Document(text=full_text, metadata=metadata)
                    return [doc]
            except Exception as e:
                logger.error(f"PyMuPDF fallback failed for '{file_path.name}': {e}")

        # 3. Ultimate Fallback
        raw_text = file_path.read_text(encoding="utf-8", errors="ignore")
        return [Document(text=raw_text, metadata=metadata)]


class DocxStructureReader(BaseReader):
    """DOCX Reader for LlamaIndex extracting paragraphs, headers, and tables formatted as Markdown grid tables."""

    def load_data(
        self,
        file: Path,
        extra_info: Optional[Dict[str, Any]] = None,
        **load_kwargs: Any,
    ) -> List[Document]:
        """Parse DOCX file into LlamaIndex Document object with table grid preservation."""
        file_path = Path(file).resolve()
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        metadata = dict(extra_info) if extra_info else {}
        metadata["total_pages"] = 1

        if docx is None:
            raw_text = file_path.read_text(encoding="utf-8", errors="ignore")
            return [Document(text=raw_text, metadata=metadata)]

        try:
            doc = docx.Document(file_path)
            content_blocks: List[str] = []

            # Extract paragraphs and preserve headers
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                style_name = p.style.name.lower() if p.style else ""
                if "heading 1" in style_name:
                    content_blocks.append(f"# {text}")
                elif "heading 2" in style_name:
                    content_blocks.append(f"## {text}")
                elif "heading 3" in style_name:
                    content_blocks.append(f"### {text}")
                elif "list" in style_name or "bullet" in style_name:
                    content_blocks.append(f"- {text}")
                else:
                    content_blocks.append(text)

            # Extract tables as clean Markdown grid tables
            for table in doc.tables:
                if not table.rows:
                    continue
                table_lines = []
                # Header row
                header_cells = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
                table_lines.append("| " + " | ".join(header_cells) + " |")
                table_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                # Data rows
                for row in table.rows[1:]:
                    row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    table_lines.append("| " + " | ".join(row_cells) + " |")
                content_blocks.append("\n".join(table_lines))

            full_text = "\n\n".join(content_blocks).strip()
            logger.info(f"Parsed DOCX '{file_path.name}': {len(full_text):,} chars extracted (Tables preserved).")
            return [Document(text=full_text, metadata=metadata)]

        except Exception as e:
            logger.error(f"Failed to parse DOCX '{file_path.name}': {e}")
            raw_text = file_path.read_text(encoding="utf-8", errors="ignore")
            return [Document(text=raw_text, metadata=metadata)]


def get_custom_file_extractors() -> Dict[str, BaseReader]:
    """Get custom LlamaIndex BaseReader instances for SimpleDirectoryReader."""
    return {
        ".pdf": StructureAwarePDFReader(),
        ".docx": DocxStructureReader(),
    }


class DocumentParser:
    """Enriches and normalizes LlamaIndex Document instances with domain metadata."""

    @staticmethod
    def infer_document_type(file_name: str) -> str:
        """Dynamically infer document_type classification based on filename patterns."""
        fn = file_name.lower()
        if "handbook" in fn:
            return "employee_handbook"
        elif "policy" in fn or "leave" in fn or "remote" in fn:
            return "policy_document"
        elif "api" in fn or "standard" in fn or "schema" in fn or "architecture" in fn:
            return "technical_standard"
        elif "guide" in fn or "sop" in fn or "manual" in fn or "workflow" in fn:
            return "operating_guide"
        elif "finance" in fn or "reimbursement" in fn or "budget" in fn or "payroll" in fn:
            return "financial_document"
        elif "security" in fn or "compliance" in fn or "audit" in fn or "incident" in fn:
            return "security_compliance"
        else:
            return "general_document"

    def enrich_document(self, document: Document) -> Document:
        """Enrich a single LlamaIndex Document instance with metadata annotations."""
        file_name = document.metadata.get("file_name", Path(document.metadata.get("file_path", "doc")).name)

        document.metadata["document_type"] = self.infer_document_type(file_name)
        if "doc_id" not in document.metadata:
            document.metadata["doc_id"] = str(uuid.uuid5(uuid.NAMESPACE_DNS, file_name))

        if "page_number" not in document.metadata:
            document.metadata["page_number"] = document.metadata.get("page_label", 1)
        if "total_pages" not in document.metadata:
            document.metadata["total_pages"] = 1

        logger.debug(f"Enriched LlamaIndex Document '{file_name}' -> type: '{document.metadata['document_type']}'")
        return document

    def enrich_all_documents(self, documents: List[Document]) -> List[Document]:
        """Enrich a list of LlamaIndex Document instances."""
        enriched = [self.enrich_document(doc) for doc in documents]
        logger.info(f"Enriched {len(enriched)} LlamaIndex Document object(s) with domain metadata.")
        return enriched


if __name__ == "__main__":
    from document_loader import DocumentLoader

    logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
    loader = DocumentLoader()
    parser = DocumentParser()

    raw_docs = loader.load_documents()
    enriched_docs = parser.enrich_all_documents(raw_docs)

    print("\n" + "=" * 80)
    print(f" ENRICHED LLAMAINDEX DOCUMENTS ({len(enriched_docs)} total)")
    print("=" * 80)
    for doc in enriched_docs[:5]:
        print(f" - [ID: {doc.doc_id[:12]}...] {doc.metadata.get('source')} | Type: {doc.metadata.get('document_type')}")
        print(f"   Snippet: {doc.text[:120]}...\n")
